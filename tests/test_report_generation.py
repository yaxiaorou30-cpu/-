import json
import shutil
import unittest
import uuid
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from src.analyzer import Analyzer
from src.orchestrator import Orchestrator, build_report_preview, generate_report
from src.preprocessor import Preprocessor
from src.report_builder import build_data_quality_summary, build_timeline, select_key_samples


class ReportGenerationTests(unittest.TestCase):
    def _sample_data(self):
        return [
            {
                "title": "警方通报突发事件处置情况",
                "content": "公安机关发布最新情况，调查处置工作正在依法推进，现场秩序已恢复。",
                "url": "https://news.example.com/a",
                "pub_time": "2026-07-02T09:12:00",
                "time_basis": "published_time",
                "source": "平安湖北",
                "platform": "百度新闻",
                "source_type": "official",
                "source_group": "stable",
                "data_type": "real",
                "collector": "百度新闻",
                "heat_index": 28.5,
                "comment_count": 12,
                "like_count": 30,
            },
            {
                "title": "网友关注事件后续进展",
                "content": "不少网民关注后续调查进展，希望权威信息及时发布，避免谣言扩散。",
                "url": "https://weibo.com/detail/1234567890",
                "pub_time": "2026-07-02T10:20:00",
                "time_basis": "published_time",
                "source": "网友甲",
                "platform": "微博",
                "source_type": "public",
                "source_group": "social",
                "data_type": "real",
                "collector": "微博搜索",
                "heat_index": 48.0,
                "comment_count": 36,
                "like_count": 120,
            },
            {
                "title": "主流媒体梳理处置进展",
                "content": "媒体对公开信息进行梳理，提醒公众以官方发布为准。",
                "url": "https://media.example.com/b",
                "pub_time": "2026-07-02",
                "time_basis": "published_time",
                "source": "中新网",
                "platform": "通用新闻搜索",
                "source_type": "media",
                "source_group": "stable",
                "data_type": "real",
                "collector": "通用新闻搜索",
                "heat_index": 18.0,
                "comment_count": 3,
                "like_count": 8,
            },
        ]

    @staticmethod
    def _twelve_catalog_records():
        return [
            {
                "title": f"完整证据目录标题 {index:02d}",
                "content": f"这是第 {index} 条完整、可追溯且已经人工审核的报告证据。",
                "url": f"https://catalog.example.com/{index}",
                "pub_time": f"2026-08-14T{index % 24:02d}:00:00",
                "time_basis": "published_time",
                "source": "目录测试来源",
                "platform": "目录测试平台",
                "source_type": "media",
                "source_group": "stable",
                "data_type": "real",
                "human_review": {
                    "reviewed_at": "2026-08-14T16:00:00",
                    "content_category": "社会事件",
                    "sentiment_label": "中性",
                },
                "content_category": "社会事件",
                "sentiment_label": "中性",
            }
            for index in range(1, 13)
        ]

    def _write_catalog_fixture(self, directory):
        root = Path(directory)
        data_path = root / "latest_news.json"
        meta_path = root / "latest_news_meta.json"
        data = self._twelve_catalog_records()
        data_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        meta_path.write_text(json.dumps({
            "topic": "完整证据目录测试",
            "keywords": ["完整证据目录"],
            "min_real_results": 12,
            "stable_sources": ["目录测试来源"],
            "review": {
                "reviewed_at": "2026-08-14T16:00:00",
                "kept_total": 12,
                "labels_confirmed": True,
            },
        }, ensure_ascii=False), encoding="utf-8")
        return data_path

    def _temporary_data_dir(self, prefix):
        root = Path("data") / f"{prefix}{uuid.uuid4().hex}"
        root.mkdir(parents=True)
        self.addCleanup(shutil.rmtree, root, True)
        return root

    def test_report_quality_and_samples_are_built_from_raw_data(self):
        data = self._sample_data()
        quality = build_data_quality_summary(data, {})
        samples = select_key_samples(data)
        timeline = build_timeline(data)
        self.assertEqual(quality["total"], 3)
        self.assertEqual(quality["stable_count"], 2)
        self.assertEqual(quality["social_count"], 1)
        self.assertEqual(quality["valid_pub_time_rate"], 1.0)
        self.assertEqual(quality["status_code"], "ready_for_review")
        self.assertTrue(quality["ready_for_review"])
        self.assertNotIn("quality_score", quality)
        self.assertEqual(samples[0]["reference_id"], "S1")
        self.assertEqual(samples[0]["platform"], "微博")
        self.assertEqual(timeline[0]["display_time"], "2026-07-02")

    def test_report_quality_counts_public_news_separately(self):
        data = [{
            "title": "公开新闻线索",
            "content": "可追溯的公开新闻摘要。",
            "url": "https://news.example.com/public/1",
            "pub_time": "2026-08-22T12:00:00",
            "time_basis": "published_time",
            "source": "示例媒体",
            "platform": "Bing 新闻",
            "source_type": "media",
            "source_group": "public_news",
            "data_type": "real",
        }]

        quality = build_data_quality_summary(data, {})

        self.assertEqual(quality["public_news_count"], 1)
        self.assertEqual(quality["statistics"]["public_news_count"], 1)

    def test_report_preview_keeps_full_evidence_catalog_beyond_eight_key_samples(self):
        Path("data").mkdir(exist_ok=True)
        tmp = self._temporary_data_dir("_test_evidence_catalog_")
        data_path = self._write_catalog_fixture(tmp)

        preview = build_report_preview(str(data_path), "event_report")

        self.assertLessEqual(len(preview["key_samples"]), 8)
        self.assertEqual(len(preview["evidence_catalog"]), 12)
        self.assertEqual(
            [item["reference_id"] for item in preview["evidence_catalog"]],
            [f"S{index}" for index in range(1, 13)],
        )
        self.assertIn(
            preview["evidence_catalog"][8]["title"],
            {record["title"] for record in self._twelve_catalog_records()},
        )

    def test_report_grounding_and_word_trace_accept_catalog_reference_beyond_s8(self):
        Path("data").mkdir(exist_ok=True)
        tmp = self._temporary_data_dir("_test_catalog_word_")
        data_path = self._write_catalog_fixture(tmp)
        output_path = Path(tmp) / "catalog-report.docx"
        preview = build_report_preview(str(data_path), "event_report")
        s9_title = preview["evidence_catalog"][8]["title"]

        generated = generate_report(
            str(data_path),
            "event_report",
            str(output_path),
            section_overrides={
                "recommendations": "继续核对第九条完整证据目录记录。[S9]",
            },
        )

        doc = Document(generated)
        all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        all_text += "\n" + "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("S9", all_text)
        self.assertIn(s9_title, all_text)

    def test_section_override_longer_than_twelve_thousand_characters_is_preserved(self):
        document_context = SimpleNamespace(
            sections={"analysis": "原内容"},
            title="",
        )
        long_override = "长篇人工分析正文。" + "甲" * 12_500 + "【尾部保留标记】"

        Orchestrator._apply_section_overrides(
            document_context,
            {"analysis": long_override},
        )

        self.assertEqual(document_context.sections["analysis"], long_override)
        self.assertTrue(document_context.sections["analysis"].endswith("【尾部保留标记】"))

    def test_timeline_skips_suspicious_old_social_dates(self):
        data = self._sample_data() + [{
            "title": "微博详情页作者资料生日不应进入时间线",
            "content": "作者资料 男 1995-09-07 处女座 湖北 查看个人主页",
            "url": "https://weibo.com/7277232894/R6FbMtX7y",
            "pub_time": "1995-09-07",
            "time_basis": "published_date",
            "source": "狂很子老",
            "platform": "微博",
            "source_type": "public",
            "source_group": "social",
            "data_type": "real",
        }]
        timeline = build_timeline(data)
        self.assertFalse(any(event["display_time"].startswith("1995") for event in timeline))

    def test_report_preview_and_docx_include_quality_blocks(self):
        tmp_path = Path("data") / "_test_report_generation"
        tmp_path.mkdir(parents=True, exist_ok=True)
        data_path = tmp_path / "latest_news.json"
        meta_path = tmp_path / "latest_news_meta.json"
        output_path = tmp_path / "report.docx"
        data_path.write_text(json.dumps(self._sample_data(), ensure_ascii=False), encoding="utf-8")
        meta_path.write_text(json.dumps({
            "topic": "测试突发事件",
            "keywords": ["测试突发事件", "警方通报"],
            "min_real_results": 3,
            "stable_sources": ["测试政府官网"],
            "social_platforms": ["微博"],
            "summary": {
                "stable_real_count": 2,
                "social_real_count": 1,
                "mock_count": 0,
            },
            "review": {
                "reviewed_at": "2026-07-31T20:00:00",
                "kept_total": 3,
                "labels_confirmed": True,
            },
        }, ensure_ascii=False), encoding="utf-8")

        preview = build_report_preview(str(data_path), "police_report")
        self.assertEqual(preview["quality"]["status_code"], "ready_for_review")
        self.assertNotIn("quality_score", preview["quality"])
        self.assertTrue(preview["sections"])
        self.assertTrue(preview["key_samples"])
        self.assertTrue(preview["quality"]["ready_for_review"])
        self.assertEqual(preview["quality"]["checks"][-1]["id"], "report_citations")
        self.assertEqual(preview["quality"]["checks"][-1]["status"], "pass")
        self.assertTrue(preview["timeline"])
        self.assertEqual(preview["analysis"]["event_keyword"], "测试突发事件")
        self.assertGreater(preview["grounding"]["citation_count"], 0)
        self.assertEqual(preview["grounding"]["unknown_sample_ids"], [])
        section_text = "\n".join(section["content"] for section in preview["sections"])
        self.assertIn("[S1]", section_text)
        self.assertIn("警方通报突发事件处置情况", section_text)
        for unsupported in (
            "已成立专项调查组",
            "涉案人员已被依法控制",
            "现场勘查已完成",
            "案件基本情况正在核查中",
        ):
            self.assertNotIn(unsupported, section_text)

        generated = generate_report(
            str(data_path),
            "police_report",
            str(output_path),
            section_overrides={"recommendations": "人工修改后的待核查建议，保留证据引用[S1]。"},
        )
        doc = Document(generated)
        paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("采集检查清单", paragraph_text)
        self.assertIn("来源覆盖表", paragraph_text)
        self.assertIn("事件时间线", paragraph_text)
        self.assertIn("重点样本与来源追溯", paragraph_text)
        self.assertIn("当前状态", table_text)
        self.assertIn("稳定源/社交平台", table_text)
        self.assertIn("报告引用", table_text)
        self.assertIn("人工修改后的待核查建议", paragraph_text)
        all_text = f"{paragraph_text}\n{table_text}"
        self.assertIn("S1", all_text)
        self.assertIn("警方通报突发事件处置情况", all_text)
        external_links = [
            rel
            for rel in doc.part.rels.values()
            if rel.reltype.endswith("/hyperlink") and rel.is_external
        ]
        self.assertGreater(len(external_links), 0)

    def test_checklist_exposes_actionable_status_instead_of_score(self):
        platforms = ["微博", "B站", "小红书", "抖音", "百度贴吧"]
        data = []
        for index in range(10):
            platform = platforms[index % len(platforms)]
            data.append({
                "title": f"菲尔兹奖线索 {index + 1}",
                "content": "这是一条可以人工核对的真实采集内容。",
                "url": f"https://example.com/{index + 1}",
                "pub_time": "2026-07-31T12:00:00" if index < 6 else "",
                "time_basis": "published_time" if index < 6 else "unknown",
                "source": platform,
                "platform": platform,
                "source_type": "public",
                "source_group": "social",
                "data_type": "real",
            })
        meta = {
            "min_real_results": 10,
            "social_platforms": platforms,
            "stable_sources": [f"政府官网 {index}" for index in range(7)],
            "failures": [],
        }

        quality = build_data_quality_summary(data, meta)
        checks = {item["id"]: item for item in quality["checks"]}

        self.assertEqual(quality["status_code"], "needs_attention")
        self.assertEqual(quality["status_label"], "需要补充或复核")
        self.assertEqual(checks["real_data"]["status"], "pass")
        self.assertEqual(checks["social_platforms"]["value"], "5/5 个已选平台有结果")
        self.assertEqual(checks["government_sources"]["status"], "warning")
        self.assertEqual(checks["publication_time"]["value"], "6/10 条已确认")
        self.assertNotIn("quality_score", quality)

    def test_report_topic_and_generic_themes_use_task_keywords(self):
        data = [
            {
                "title": "一口气看懂王虹获得菲尔兹奖：挂谷猜想为什么重要",
                "content": "介绍王虹的数学研究、挂谷猜想和科研背景。",
                "url": "https://example.com/math",
                "pub_time": "2026-07-31T12:00:00",
                "time_basis": "published_time",
                "source": "科普作者",
                "platform": "B站",
                "source_type": "public",
                "source_group": "social",
                "data_type": "real",
            },
            {
                "title": "从北大学习经历到海外博士阶段",
                "content": "文章讨论人物经历、大学教育和科研履历。",
                "url": "https://example.com/person",
                "pub_time": "2026-07-31T13:00:00",
                "time_basis": "published_time",
                "source": "平台用户",
                "platform": "微博",
                "source_type": "public",
                "source_group": "social",
                "data_type": "real",
            },
            {
                "title": "祝贺获奖并讨论研究意义",
                "content": "网友表达祝贺，也评价这项研究的重要性。",
                "url": "https://example.com/opinion",
                "pub_time": "2026-07-31T14:00:00",
                "time_basis": "published_time",
                "source": "平台用户乙",
                "platform": "小红书",
                "source_type": "public",
                "source_group": "social",
                "data_type": "real",
            },
        ]
        tmp_path = Path("data") / "_test_report_topics"
        tmp_path.mkdir(parents=True, exist_ok=True)
        data_path = tmp_path / "topic.json"
        meta_path = tmp_path / "topic_meta.json"
        data_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        meta_path.write_text(json.dumps({
            "topic": "菲尔兹奖2026.7.31",
            "keywords": ["菲尔兹奖", "王虹"],
            "min_real_results": 3,
            "social_platforms": ["B站", "微博", "小红书"],
            "review": {"reviewed_at": "2026-07-31T20:00:00", "kept_total": 3, "labels_confirmed": True},
        }, ensure_ascii=False), encoding="utf-8")

        preview = build_report_preview(str(data_path), "event_report")
        section_text = "\n".join(section["content"] for section in preview["sections"])
        analysis_text = next(
            section["content"]
            for section in preview["sections"]
            if section["id"] == "analysis"
        )

        self.assertEqual(preview["analysis"]["event_keyword"], "菲尔兹奖")
        self.assertEqual(preview["analysis"]["query_keywords"], ["菲尔兹奖", "王虹"])
        self.assertEqual(preview["analysis"]["top_keywords"][:2], ["菲尔兹奖", "王虹"])
        self.assertIn("专业解释与背景", section_text)
        self.assertIn("人物经历与机构关联", section_text)
        self.assertIn("代表内容线索（每条仅列一次）", analysis_text)
        for sample in preview["key_samples"]:
            self.assertLessEqual(
                analysis_text.count(sample["title"]),
                1,
                f"同一内容线索在主要观点与研判中重复：{sample['reference_id']}",
            )
        self.assertGreaterEqual(analysis_text.count("[S1]"), 2)
        self.assertNotIn("灾情数字、人员情况、停课停运", section_text)
        self.assertEqual(
            {item["id"]: item for item in preview["quality"]["checks"]}["template_fit"]["status"],
            "pass",
        )

    def test_report_generation_rejects_unknown_sample_reference(self):
        tmp_path = Path("data") / "_test_report_generation"
        tmp_path.mkdir(parents=True, exist_ok=True)
        data_path = tmp_path / "latest_news.json"
        meta_path = tmp_path / "latest_news_meta.json"
        output_path = tmp_path / "unknown-reference.docx"
        data_path.write_text(json.dumps(self._sample_data(), ensure_ascii=False), encoding="utf-8")
        meta_path.write_text(json.dumps({
            "topic": "测试案件",
            "keywords": ["测试案件"],
        }, ensure_ascii=False), encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "未知样本编号"):
            generate_report(
                str(data_path),
                "police_report",
                str(output_path),
                section_overrides={"recommendations": "错误引用[S999]。"},
            )

    def test_missing_pub_time_does_not_extend_analysis_time_range(self):
        data = self._sample_data()[:1] + [{
            "title": "没有可靠发布时间的社交平台内容",
            "content": "该条记录只有采集时间，不应把采集时间写成内容发布时间。",
            "url": "https://example.com/no-time",
            "pub_time": "",
            "time_basis": "unknown",
            "source": "测试账号",
            "platform": "微博",
            "source_type": "public",
            "source_group": "social",
            "data_type": "real",
        }]
        records = Preprocessor().process(data)
        context = Analyzer().analyze(records)
        self.assertEqual(context.confirmed_time_count, 1)
        self.assertEqual(context.time_range[0], context.time_range[1])
        self.assertEqual(context.time_range[0].strftime("%Y-%m-%d"), "2026-07-02")

    def test_report_preview_rejects_failed_body_until_original_is_checked(self):
        root = self._temporary_data_dir("_test_pending_body_")
        data_path = root / "latest_news.json"
        data_path.write_text(
            json.dumps([{
                **self._sample_data()[0],
                "body_fetch_status": "failed",
            }], ensure_ascii=False),
            encoding="utf-8",
        )

        with self.assertRaisesRegex(ValueError, "正文获取失败"):
            build_report_preview(str(data_path), "event_report")

    def test_multi_event_reports_reject_failed_body_until_original_is_checked(self):
        root = self._temporary_data_dir("_test_pending_body_multi_")
        data_path = root / "latest_news.json"
        data_path.write_text(
            json.dumps([{
                **self._sample_data()[0],
                "body_fetch_status": "failed",
            }], ensure_ascii=False),
            encoding="utf-8",
        )

        with self.assertRaisesRegex(ValueError, "正文获取失败"):
            Orchestrator().generate_multi_event_reports(
                str(data_path),
                str(root / "reports"),
            )

    def test_generic_social_content_does_not_create_case_progress(self):
        records = Preprocessor().process([{
            "title": "网友讨论一则社会事件",
            "content": "网友正在讨论事件经过，当前采集内容没有任何案件处置进展。",
            "url": "https://example.com/social",
            "pub_time": "",
            "time_basis": "unknown",
            "source": "网友",
            "platform": "微博",
            "source_type": "public",
            "source_group": "social",
            "data_type": "real",
        }])
        context = Analyzer().analyze(records)
        self.assertEqual(context.case_progress, "")


if __name__ == "__main__":
    unittest.main()
