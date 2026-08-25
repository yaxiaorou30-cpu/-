import os
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from src.models import DocumentContext
from src.utils.logger import get_logger

logger = get_logger(__name__)


class Renderer:
    def __init__(self):
        pass

    def render(self, document_context: DocumentContext, docx_template_path: Optional[str], output_path: str):
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)

        if docx_template_path and os.path.exists(docx_template_path):
            self._render_with_template(document_context, docx_template_path, output_path)
        else:
            self._render_from_scratch(document_context, output_path)

        logger.info(f"文档已渲染: {output_path}")

    def _render_with_template(self, ctx: DocumentContext, template_path: str, output_path: str):
        doc = Document(template_path)
        bookmarks = self._extract_bookmarks(doc)

        for section_id, text in ctx.sections.items():
            bookmark_name = f"bookmark_{section_id}"
            if bookmark_name in bookmarks:
                self._replace_bookmark_text(bookmarks[bookmark_name], text)
            else:
                # 如果没有对应书签，在文档末尾追加
                self._append_section(doc, section_id, text)

        self._append_report_quality(doc, ctx)
        self._append_source_distribution(doc, ctx)
        self._append_timeline(doc, ctx)
        self._append_key_samples(doc, ctx)

        doc.save(output_path)

    def _render_from_scratch(self, ctx: DocumentContext, output_path: str):
        doc = Document()
        self._setup_document(doc)

        chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]

        section_order = ctx.metadata.get("section_order", list(ctx.sections.keys()))
        section_names = ctx.metadata.get("section_names", {})
        section_review = ctx.metadata.get("section_review", {})

        section_idx = 0
        for section_id in section_order:
            if section_id not in ctx.sections:
                continue

            text = ctx.sections[section_id]

            if section_id == "title":
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._append_generated_note(doc, ctx)
            else:
                section_name = section_names.get(section_id, section_id)
                if section_name:
                    chinese_title = f"{chinese_nums[section_idx]}、{section_name}"
                    doc.add_heading(chinese_title, level=1)
                    section_idx += 1
                else:
                    doc.add_heading(section_id, level=1)

                self._append_body_text(doc, text)

                if section_review.get(section_id, False):
                    review_p = doc.add_paragraph()
                    run = review_p.add_run("（建议人工审定本章内容）")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.font.italic = True
                    review_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self._append_report_quality(doc, ctx)
        self._append_source_distribution(doc, ctx)
        self._append_timeline(doc, ctx)
        self._append_key_samples(doc, ctx)

        doc.save(output_path)

    @staticmethod
    def _extract_bookmarks(doc):
        """提取文档中的书签映射。python-docx 书签访问较底层，这里做简化处理。"""
        # python-docx 没有原生书签遍历 API，此函数预留扩展
        return {}

    @staticmethod
    def _replace_bookmark_text(element, text: str):
        # 预留：通过底层 XML 操作替换书签范围文本
        pass

    @staticmethod
    def _append_section(doc: Document, section_id: str, text: str):
        if section_id == "title":
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_heading(section_id, level=1)
            doc.add_paragraph(text)

    @staticmethod
    def _setup_document(doc: Document):
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        for style_name in ("Heading 1", "Heading 2", "Title"):
            style = doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    @staticmethod
    def _append_generated_note(doc: Document, ctx: DocumentContext):
        quality = ctx.metadata.get("report_quality") or {}
        analysis = ctx.metadata.get("analysis_summary") or {}
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_at = quality.get("generated_at") or ctx.metadata.get("generated_at", "")
        generated_at = str(generated_at).replace("T", " ")
        event_keyword = analysis.get("event_keyword") or ctx.event_keyword or "舆情事件"
        run = note.add_run(f"任务主题：{event_keyword}    生成时间：{generated_at}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 112, 133)

    def _append_report_quality(self, doc: Document, ctx: DocumentContext):
        quality = ctx.metadata.get("report_quality") or {}
        if not quality:
            return
        doc.add_heading("采集检查清单", level=1)
        rows = [
            ("当前状态", quality.get("status_label", "未检查")),
            ("状态说明", quality.get("status_detail", "")),
            ("样本总量", quality.get("total", 0)),
            ("真实/模拟", f"{quality.get('real_count', 0)} / {quality.get('mock_count', 0)}"),
            ("稳定源/社交平台", f"{quality.get('stable_count', 0)} / {quality.get('social_count', 0)}"),
        ]
        self._add_key_value_table(doc, rows)

        checks = quality.get("checks") or []
        if checks:
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for index, label in enumerate(("检查项", "状态", "实际结果", "说明")):
                cell = table.rows[0].cells[index]
                cell.text = label
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for item in checks:
                cells = table.add_row().cells
                cells[0].text = str(item.get("label") or "-")
                cells[1].text = str(item.get("status_label") or "-")
                cells[2].text = str(item.get("value") or "-")
                cells[3].text = str(item.get("detail") or "-")

    def _append_source_distribution(self, doc: Document, ctx: DocumentContext):
        quality = ctx.metadata.get("report_quality") or {}
        total = int(quality.get("total") or 0)
        platform_dist = quality.get("platform_distribution") or {}
        source_type_dist = quality.get("source_type_distribution") or {}
        if not total or (not platform_dist and not source_type_dist):
            return

        doc.add_heading("来源覆盖表", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["维度", "项目", "数量", "占比"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        def add_rows(label, dist):
            for name, count in sorted(dist.items(), key=lambda item: item[1], reverse=True):
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = str(name or "未知")
                cells[2].text = str(count)
                cells[3].text = f"{(count / total):.0%}" if total else "0%"

        add_rows("采集渠道", platform_dist)
        add_rows("来源类型", source_type_dist)

    def _append_timeline(self, doc: Document, ctx: DocumentContext):
        timeline = ctx.metadata.get("timeline") or []
        if not timeline:
            return
        doc.add_heading("事件时间线", level=1)
        for event in timeline:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            label = paragraph.add_run(
                f"{event.get('display_time') or event.get('time') or '-'}  "
                f"{event.get('platform') or '-'} · {event.get('source') or '-'}"
            )
            label.bold = True
            paragraph.add_run(f"\n{event.get('title') or '(无标题)'}  ")
            url = str(event.get("url") or "")
            if url.startswith(("http://", "https://")):
                self._append_hyperlink(paragraph, url, "打开原文")

    def _append_key_samples(self, doc: Document, ctx: DocumentContext):
        samples = ctx.metadata.get("key_samples") or []
        if not samples:
            return
        section_heading = doc.add_heading("重点样本与来源追溯", level=1)
        section_heading.paragraph_format.keep_with_next = True
        note = doc.add_paragraph("引用编号用于把报告结论与原始采集记录对应；来源链接来自 data/latest_news.json。")
        note.paragraph_format.space_after = Pt(6)
        note.paragraph_format.keep_with_next = True
        for index, sample in enumerate(samples, 1):
            reference_id = sample.get("reference_id") or f"S{index}"
            heading = doc.add_heading(
                f"{reference_id} · {sample.get('platform') or '-'} · {sample.get('source') or '-'}",
                level=2,
            )
            heading.paragraph_format.keep_with_next = True

            meta = doc.add_paragraph()
            meta.paragraph_format.space_after = Pt(3)
            meta.paragraph_format.keep_with_next = True
            meta.add_run("发布时间：").bold = True
            meta.add_run(str(sample.get("pub_time") or "时间未知"))

            analysis = doc.add_paragraph()
            analysis.paragraph_format.space_after = Pt(3)
            analysis.paragraph_format.keep_with_next = True
            analysis.add_run("审核标签：").bold = True
            analysis.add_run(
                f"{sample.get('content_category') or '其他'}；"
                f"情感参考：{sample.get('sentiment_label') or '中性'}"
            )
            if sample.get("review_note"):
                analysis.add_run(f"；人工备注：{sample.get('review_note')}")

            title = doc.add_paragraph()
            title.paragraph_format.space_after = Pt(3)
            title.paragraph_format.keep_with_next = True
            title.add_run("标题：").bold = True
            title.add_run(str(sample.get("title") or "(无标题)"))

            excerpt = doc.add_paragraph()
            excerpt.paragraph_format.space_after = Pt(3)
            excerpt.paragraph_format.keep_with_next = True
            excerpt.add_run("内容摘要：").bold = True
            excerpt.add_run(str(sample.get("content_excerpt") or "(无正文摘要)"))

            link = doc.add_paragraph()
            link.paragraph_format.space_after = Pt(8)
            link.add_run("来源链接：").bold = True
            url = str(sample.get("url") or "")
            if url.startswith(("http://", "https://")):
                self._append_hyperlink(link, url, "打开原文")
            else:
                link.add_run("无有效链接")

    @staticmethod
    def _add_key_value_table(doc: Document, rows):
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for key, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)
            if cells[0].paragraphs and cells[0].paragraphs[0].runs:
                cells[0].paragraphs[0].runs[0].bold = True

    @staticmethod
    def _set_cell_hyperlink(cell, url: str, display_text: str):
        cell.text = ""
        Renderer._append_hyperlink(cell.paragraphs[0], url, display_text)

    @staticmethod
    def _append_hyperlink(paragraph, url: str, display_text: str):
        relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relation_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(color)
        run_properties.append(underline)
        run.append(run_properties)

        text = OxmlElement("w:t")
        text.text = display_text
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _append_body_text(doc: Document, text: str):
        parts = [part.strip() for part in str(text or "").split("\n") if part.strip()]
        if not parts:
            doc.add_paragraph("")
            return
        for part in parts:
            paragraph = doc.add_paragraph(part)
            paragraph.paragraph_format.first_line_indent = Pt(22)
            paragraph.paragraph_format.space_after = Pt(6)
